import json
import os
import re
import time
from pathlib import Path
from typing import Any, Dict, List

import pandas as pd
import streamlit as st
from docx import Document as DocxDocument
from pypdf import PdfReader

try:
    from google import genai as google_genai
except Exception:
    google_genai = None

CHAT_FILE = Path("chat_data.json")
PROJECT_TITLE = "AI Powered Multi Doc Conversational Chatbot Using RAG"
SUPPORTED_TYPES = ["pdf", "txt", "xml", "docx", "json", "csv"]
ALLOWED_ANSWER_STYLES = ["Detailed", "Short", "Bullet Points", "Professional"]
MAX_CHUNK_SIZE = 900
CHUNK_OVERLAP = 120
MAX_CONTEXT_CHUNKS = 6
MAX_API_RETRIES = 3
API_KEY_ERROR_PATTERNS = ["API_KEY_INVALID", "key expired", "api key expired", "invalid api key"]
MAX_CHUNKS_PER_FILE = 2


def apply_ui_styles() -> None:
    st.markdown(
        """
<style>
:root {
    --app-bg: #08070d;
    --panel-bg: #171420;
    --panel-2: #211d2d;
    --panel-border: #5b5472;
    --primary: #7df7d5;
    --primary-2: #c88bff;
    --accent: #ff9870;
    --text-main: #f5f2ff;
    --muted: #bab3cc;
    --success: #4ade80;
    --danger: #f87171;
}
[data-testid="stAppViewContainer"] {
    background: transparent;
}
[data-testid="stAppViewContainer"] .main {
    background: transparent;
}
[data-testid="stHeader"] {
    background: rgba(13, 10, 22, 0.62);
    border-bottom: 1px solid #4f4a63;
}
.stApp {
    background:
        radial-gradient(1200px 700px at -8% -12%, #d47eff2a 0%, transparent 68%),
        radial-gradient(1100px 520px at 120% 8%, #58f0c82a 0%, transparent 65%),
        radial-gradient(900px 480px at 45% 115%, #ff8f7030 0%, transparent 70%),
        linear-gradient(118deg, #08070d 0%, #12101a 45%, #151422 100%);
    background-size: 130% 130%, 140% 140%, 150% 150%;
    animation: bgShift 14s ease-in-out infinite;
    color: var(--text-main);
}
div[data-testid="stMetric"] {
    background: linear-gradient(145deg, #1b1826 0%, #242131 100%);
    border: 1px solid var(--panel-border);
    border-radius: 14px;
    padding: 10px 12px;
    box-shadow: inset 0 1px 0 #ffffff0a, 0 8px 20px #02071166;
    animation: riseIn 0.45s ease-out 1;
}
div[data-testid="stMetricLabel"] p {
    font-size: 0.8rem;
    font-weight: 700;
    color: var(--muted);
}
div[data-testid="stMetricValue"] {
    font-size: 1.4rem;
    color: var(--text-main);
}
.block-container {
    padding-top: 1.6rem;
}
.app-header {
    background: linear-gradient(132deg, #241934 0%, #302245 56%, #223532 100%);
    background-size: 170% 170%;
    border: 1px solid #6e618b;
    border-radius: 18px;
    padding: 24px 28px;
    margin-bottom: 0.6rem;
    box-shadow: 0 14px 34px #02071488;
    animation: gradientShift 10s ease infinite, edgePulse 4s ease-in-out infinite;
    position: relative;
    overflow: hidden;
}
.app-header::after {
    content: "";
    position: absolute;
    top: 0;
    left: -40%;
    width: 40%;
    height: 100%;
    background: linear-gradient(90deg, transparent 0%, #ffffff14 50%, transparent 100%);
    transform: skewX(-20deg);
    animation: sheen 7s linear infinite;
}
.app-header h1 {
    margin: 0;
    font-size: clamp(2rem, 2.5vw, 2.7rem);
    color: #f5f2ff;
    background: linear-gradient(90deg, #f6f2ff 0%, #b996ff 46%, #7df7d5 100%);
    background-size: 180% 180%;
    -webkit-background-clip: text;
    background-clip: text;
    -webkit-text-fill-color: transparent;
    letter-spacing: 0.3px;
    animation: titleFlow 8s ease-in-out infinite;
}
.app-header p {
    margin: 0.45rem 0 0 0;
    font-size: 1.05rem;
    color: var(--muted);
}
.upload-panel {
    background: linear-gradient(180deg, #1d1828 0%, #292134 64%, #23322f 100%);
    border: 1px solid #645c7c;
    border-radius: 16px;
    padding: 14px 16px;
    margin-bottom: 0.7rem;
    animation: riseIn 0.4s ease-out 1;
    box-shadow: inset 0 1px 0 #ffffff0a, 0 10px 24px #02071166;
}
.stChatMessage {
    background: linear-gradient(180deg, #1b1825 0%, #232032 100%);
    border: 1px solid #61597a;
    border-radius: 16px;
    animation: riseIn 0.22s ease-out 1;
    transition: transform 0.2s ease, border-color 0.2s ease, box-shadow 0.2s ease;
}
.stChatMessage:hover {
    transform: translateY(-2px);
    border-color: #9ff5dc;
    box-shadow: 0 10px 22px #19122a66;
}
.stChatInput > div {
    background: #1e1a2a;
    border-radius: 20px;
    border: 1px solid #6f6288;
    box-shadow: inset 0 1px 0 #ffffff0a;
}
div[data-testid="stChatInput"] textarea:focus {
    box-shadow: 0 0 0 1px #9ff5dc88;
}
.stCaption {
    color: var(--muted);
}
[data-testid="stMarkdownContainer"] code {
    background: #191522;
    color: #f3e9ff;
}
div[data-testid="stCodeBlock"] {
    border: 1px solid #645d7c;
    border-radius: 12px;
}
section[data-testid="stFileUploader"] {
    background: linear-gradient(135deg, #211b2c 0%, #262033 100%);
    border: 1.5px dashed #b58ef0;
    border-radius: 14px;
    padding: 8px;
    animation: dashPulse 2.8s ease-in-out infinite;
}
section[data-testid="stFileUploader"] button {
    border-radius: 10px;
}
.stButton > button {
    border: 1px solid #6f608d;
    background: linear-gradient(180deg, #2b223d 0%, #251e33 100%);
    color: #f5f2ff;
    background-size: 160% 160%;
    animation: buttonFlow 6s ease-in-out infinite;
    transition: all 0.22s ease;
}
.stButton > button:hover {
    border-color: #a7f7df;
    box-shadow: 0 0 0 1px #a7f7df66, 0 8px 22px #9a6aff33;
    transform: translateY(-2px);
}
.pipeline-board {
    border: 1px solid #625b7a;
    border-radius: 14px;
    padding: 10px 10px 8px 10px;
    background: linear-gradient(180deg, #1c1727 0%, #171321 100%);
    margin-bottom: 0.8rem;
}
.pipeline-step {
    border: 1px solid #5b5470;
    border-radius: 10px;
    padding: 8px 10px;
    margin: 7px 0;
    background: #1d1828;
}
.pipeline-step strong {
    color: #f3eeff;
}
.badge-ok {
    color: #89f7b0;
}
.badge-wait {
    color: #f3d27d;
}
.badge-warn {
    color: #ff9e9e;
}
.pipeline-kpi {
    border: 1px solid #5b5470;
    border-radius: 10px;
    padding: 8px 10px;
    margin-top: 8px;
    background: #1d1828;
}
.rag-lite {
    border: 1px solid #5b5470;
    border-radius: 12px;
    padding: 10px;
    background: #1d1828;
    margin-bottom: 0.7rem;
}
.rag-lite-head {
    color: #eef5ff;
    font-weight: 700;
    font-size: 0.92rem;
    margin-bottom: 6px;
}
.rag-lite-sub {
    color: #b8cae6;
    font-size: 0.8rem;
    margin-bottom: 8px;
}
.rag-flow-row {
    display: flex;
    justify-content: space-between;
    align-items: center;
    gap: 8px;
    padding: 6px 0;
    border-bottom: 1px dashed #3e4f72;
}
.rag-flow-row:last-child {
    border-bottom: none;
}
.rag-stage {
    color: #e6f0ff;
    font-size: 0.84rem;
}
.rag-pill {
    border-radius: 999px;
    padding: 2px 8px;
    font-size: 0.72rem;
    font-weight: 700;
    letter-spacing: 0.2px;
}
.rag-pill.ok {
    background: #133424;
    color: #8ff0bc;
    border: 1px solid #2b7a53;
}
.rag-pill.wait {
    background: #2e2a16;
    color: #f1d88b;
    border: 1px solid #736535;
}
.rag-pill.off {
    background: #321919;
    color: #ffb4b4;
    border: 1px solid #7e3b3b;
}
.rag-progress {
    width: 100%;
    height: 8px;
    border-radius: 999px;
    background: #101a30;
    border: 1px solid #37537f;
    overflow: hidden;
    margin-top: 10px;
}
.rag-progress > div {
    height: 100%;
    background: linear-gradient(90deg, #58b8ff 0%, #49ded0 100%);
}
.rag-mini {
    margin-top: 8px;
    color: #b8cae6;
    font-size: 0.76rem;
}
.runtime-status {
    border: 1px solid #675c82;
    border-radius: 12px;
    padding: 8px 10px;
    margin-bottom: 10px;
    background: linear-gradient(180deg, #2a2338 0%, #312942 100%);
    display: flex;
    align-items: center;
    gap: 8px;
    font-size: 0.92rem;
}
.runtime-status strong {
    color: #f3eeff;
}
.runtime-status span {
    color: #beb6d0;
}
.runtime-dot {
    width: 8px;
    height: 8px;
    border-radius: 50%;
    background: #7df7d5;
    box-shadow: 0 0 0 0 #7df7d588;
    animation: statusPulse 1.35s ease-in-out infinite;
    flex-shrink: 0;
}
.runtime-status.done {
    border-color: #3f7b63;
    background: linear-gradient(180deg, #122521 0%, #17322b 100%);
}
.runtime-status.done .runtime-dot {
    background: #52d68f;
    animation: none;
    box-shadow: none;
}
hr {
    border: none;
    height: 2px;
    border-radius: 999px;
    background: linear-gradient(90deg, transparent 0%, #c88bff 45%, #7df7d5 55%, transparent 100%);
    background-size: 200% 200%;
    animation: barRun 6s linear infinite;
}
@keyframes gradientShift {
    0% { background-position: 0% 50%; }
    50% { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
}
@keyframes bgShift {
    0% { background-position: 0% 50%, 100% 0%, 0% 50%; }
    50% { background-position: 100% 50%, 0% 100%, 100% 50%; }
    100% { background-position: 0% 50%, 100% 0%, 0% 50%; }
}
@keyframes orbOne {
    0% { transform: translate(0, 0); }
    50% { transform: translate(45px, 30px); }
    100% { transform: translate(0, 0); }
}
@keyframes orbTwo {
    0% { transform: translate(0, 0); }
    50% { transform: translate(-40px, -26px); }
    100% { transform: translate(0, 0); }
}
@keyframes edgePulse {
    0% { box-shadow: 0 12px 28px #03091688, 0 0 0 0 #5cb9ff00; }
    50% { box-shadow: 0 12px 28px #03091688, 0 0 0 1px #5cb9ff44; }
    100% { box-shadow: 0 12px 28px #03091688, 0 0 0 0 #5cb9ff00; }
}
@keyframes titleFlow {
    0% { background-position: 0% 50%; }
    50% { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
}
@keyframes sheen {
    0% { left: -45%; }
    100% { left: 140%; }
}
@keyframes riseIn {
    from { opacity: 0.45; transform: translateY(8px); }
    to { opacity: 1; transform: translateY(0); }
}
@keyframes dashPulse {
    0% { border-color: #8a6cc7; }
    50% { border-color: #d3b7ff; }
    100% { border-color: #8a6cc7; }
}
@keyframes buttonFlow {
    0% { background-position: 0% 50%; }
    50% { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
}
@keyframes barRun {
    0% { background-position: 0% 50%; }
    100% { background-position: 200% 50%; }
}
@keyframes statusPulse {
    0% { box-shadow: 0 0 0 0 #7df7d588; }
    70% { box-shadow: 0 0 0 8px #7df7d500; }
    100% { box-shadow: 0 0 0 0 #7df7d500; }
}
@media (prefers-reduced-motion: reduce) {
    * {
        animation: none !important;
        transition: none !important;
    }
}
@media (max-width: 900px) {
    .app-header {
        padding: 16px 14px;
        border-radius: 14px;
    }
    .app-header h1 {
        font-size: 1.55rem;
        line-height: 1.25;
        letter-spacing: 0.15px;
    }
    .app-header p {
        font-size: 0.95rem;
    }
    .upload-panel {
        padding: 11px 10px;
    }
    .stChatMessage {
        border-radius: 12px;
    }
}
/* Blue Readable Theme Override */
:root {
    --app-bg: #07101f !important;
    --panel-bg: #13243f !important;
    --panel-2: #193154 !important;
    --panel-border: #3f6699 !important;
    --text-main: #eef5ff !important;
    --muted: #b8cae6 !important;
}
[data-testid="stHeader"] {
    background: #07101f !important;
    border-bottom: 1px solid #2f4f7c !important;
}
section[data-testid="stSidebar"] {
    background: #07101f !important;
}
.stApp {
    background: linear-gradient(140deg, #07101f 0%, #0b1b35 48%, #0f2445 100%) !important;
    animation: none !important;
}
.app-header {
    background: #172a4d !important;
    border: 1px solid #4670a8 !important;
}
.app-header h1 {
    color: #eef5ff !important;
    background: none !important;
    -webkit-background-clip: initial !important;
    background-clip: initial !important;
    -webkit-text-fill-color: #eef5ff !important;
    animation: none !important;
}
.upload-panel,
.stChatMessage,
.pipeline-board,
.pipeline-step,
.pipeline-kpi,
.runtime-status {
    background: #13243f !important;
    border-color: #3f6699 !important;
}
.stChatInput > div,
section[data-testid="stFileUploader"] {
    background: #10203a !important;
    border-color: #3f6699 !important;
}
.stButton > button {
    background: #1a3258 !important;
    border-color: #5382bf !important;
    color: #eef5ff !important;
}
.stButton > button:hover {
    border-color: #89c8ff !important;
    box-shadow: 0 0 0 1px #89c8ff66, 0 8px 18px #2e6ab533 !important;
}
.pipeline-step strong,
.runtime-status strong {
    color: #e8f2ff !important;
}
.runtime-status span,
.stCaption {
    color: #b8cae6 !important;
}
hr {
    background: #4670a8 !important;
}
/* Ensure blue skin on specific requested sections */
div[data-testid="stSelectbox"] > div,
div[data-testid="stSelectbox"] div[data-baseweb="select"] > div {
    background: #1a3258 !important;
    border-color: #5382bf !important;
    color: #eef5ff !important;
}
section[data-testid="stSidebar"] div[data-testid="stMetric"] {
    background: #13243f !important;
    border: 1px solid #3f6699 !important;
}
section[data-testid="stSidebar"] div[data-testid="stMetricLabel"] p,
section[data-testid="stSidebar"] div[data-testid="stMetricValue"] {
    color: #e8f2ff !important;
}
section[data-testid="stFileUploader"] > div {
    background: #10203a !important;
    border: 1.5px dashed #65a3ff !important;
}
section[data-testid="stFileUploaderDropzone"] {
    background: #10203a !important;
}
.rag-lite,
.rag-flow-row,
.rag-progress,
.rag-mini {
    background: #13243f !important;
    border-color: #3f6699 !important;
    color: #e8f2ff !important;
}
.rag-stage,
.rag-lite-head,
.rag-lite-sub {
    color: #e8f2ff !important;
}
.rag-pill.ok {
    background: #133d2f !important;
    color: #9af4cd !important;
    border-color: #2f7f64 !important;
}
.rag-pill.wait {
    background: #2f2a18 !important;
    color: #efd98e !important;
    border-color: #72673a !important;
}
.rag-pill.off {
    background: #3b1c24 !important;
    color: #ffb0bf !important;
    border-color: #824154 !important;
}
.rag-progress > div {
    background: linear-gradient(90deg, #58b8ff 0%, #49ded0 100%) !important;
}
</style>
""",
        unsafe_allow_html=True,
    )


@st.cache_resource(show_spinner=False)
def configure_model(api_key: str) -> Dict[str, Any] | None:
    if not api_key:
        return None

    if google_genai is not None:
        try:
            client = google_genai.Client(api_key=api_key)
            return {
                "backend": "google_genai",
                "client": client,
                "model_name": "gemini-2.5-flash",
            }
        except Exception:
            pass

    try:
        # Legacy fallback only if new SDK is unavailable.
        import importlib

        legacy_genai = importlib.import_module("google.generativeai")
        legacy_genai.configure(api_key=api_key)
        legacy_model = legacy_genai.GenerativeModel("models/gemini-1.5-flash")
        return {
            "backend": "legacy",
            "model": legacy_model,
            "model_name": "models/gemini-1.5-flash",
        }
    except Exception:
        return None


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def split_text(text: str, size: int = MAX_CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    cleaned = normalize_spaces(text)
    if not cleaned:
        return []

    if size <= overlap:
        overlap = 0

    chunks: List[str] = []
    step = size - overlap
    start = 0
    while start < len(cleaned):
        end = min(start + size, len(cleaned))
        if end < len(cleaned):
            boundary = cleaned.rfind(" ", start, end)
            if boundary != -1 and boundary > start + int(size * 0.55):
                end = boundary
        chunk = cleaned[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(cleaned):
            break
        start = max(end - overlap, start + 1)
    return chunks


def chunkify(text: str, file_name: str, location: str) -> List[Dict]:
    chunks = split_text(text)
    output: List[Dict] = []
    for idx, chunk in enumerate(chunks, start=1):
        output.append(
            {
                "chunk_id": f"{file_name}|{location}|{idx}",
                "file_name": file_name,
                "location": location,
                "text": chunk,
                "preview": chunk[:260],
            }
        )
    return output


def extract_structure_lines(text: str) -> str:
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        return ""

    structure_keywords = ("step", "process", "decision", "start", "end", "input", "output", "flow")
    matched: List[str] = []
    for line in lines:
        lower = line.lower()
        is_table_like = "|" in line or "\t" in line or re.search(r"\s{2,}", line)
        is_flow_like = "->" in line or "=>" in line or any(key in lower for key in structure_keywords)
        if is_table_like or is_flow_like:
            matched.append(line)
    return "\n".join(matched[:120])


def extract_docx_table_text(doc: DocxDocument) -> str:
    blocks: List[str] = []
    for t_index, table in enumerate(doc.tables, start=1):
        rows: List[str] = []
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", cell.text or "").strip() for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(f"Table {t_index}:\n" + "\n".join(rows))
    return "\n\n".join(blocks)


def load_uploaded_file(file) -> List[Dict]:
    name = file.name.lower()
    records: List[Dict] = []

    try:
        if name.endswith(".pdf"):
            reader = PdfReader(file)
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text(extraction_mode="layout") or page.extract_text() or ""
                records.extend(chunkify(page_text, file.name, f"Page {page_num}"))
                structured = extract_structure_lines(page_text)
                if structured:
                    records.extend(chunkify(structured, file.name, f"Page {page_num} Structured"))
            return records

        if name.endswith((".txt", ".xml")):
            content = file.read().decode("utf-8", "ignore")
            chunks = chunkify(content, file.name, "Section 1")
            structured = extract_structure_lines(content)
            if structured:
                chunks.extend(chunkify(structured, file.name, "Structured Lines"))
            return chunks

        if name.endswith(".docx"):
            doc = DocxDocument(file)
            paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = extract_docx_table_text(doc)
            full_text = "\n\n".join(part for part in [paragraph_text, table_text] if part.strip())
            chunks = chunkify(full_text, file.name, "Section 1")
            structured = extract_structure_lines(full_text)
            if structured:
                chunks.extend(chunkify(structured, file.name, "Structured Lines"))
            return chunks

        if name.endswith(".json"):
            parsed = json.load(file)
            content = json.dumps(parsed, indent=2, ensure_ascii=False)
            return chunkify(content, file.name, "Section 1")

        if name.endswith(".csv"):
            df = pd.read_csv(file)
            text = df.to_string(index=False)
            chunks = chunkify(text, file.name, "Table")
            structured = extract_structure_lines(text)
            if structured:
                chunks.extend(chunkify(structured, file.name, "Table Structured"))
            return chunks

        st.warning(f"Unsupported file type: {file.name}")
        return []
    except Exception as exc:
        st.error(f"Could not read {file.name}: {exc}")
        return []


def save_chat(messages: List[Dict]) -> None:
    try:
        CHAT_FILE.write_text(json.dumps(messages, ensure_ascii=False, indent=2), encoding="utf-8")
    except OSError:
        st.warning("Could not persist chat history to disk.")


def load_chat() -> List[Dict]:
    if not CHAT_FILE.exists():
        return []
    try:
        data = json.loads(CHAT_FILE.read_text(encoding="utf-8"))
        if isinstance(data, list):
            for item in data:
                content = str(item.get("content", ""))
                if is_api_key_issue(content):
                    item["content"] = (
                        "API key issue detected earlier. Update GEMINI_API_KEY to restore live API responses."
                    )
            return data
    except (json.JSONDecodeError, OSError):
        pass
    return []


def score_chunk(query: str, chunk: str) -> int:
    terms = [word.strip(".,!?;:\"'()[]{}").lower() for word in query.split()]
    terms = [word for word in terms if len(word) > 2]
    if not terms:
        return 0

    chunk_lower = chunk.lower()
    overlap_score = sum(2 for word in terms if word in chunk_lower)
    phrase_bonus = 4 if " ".join(terms[:3]) in chunk_lower else 0
    return overlap_score + phrase_bonus


def top_relevant_chunks(query: str, data: List[Dict], limit: int = MAX_CONTEXT_CHUNKS) -> List[Dict]:
    def normalize_file_token(name: str) -> str:
        base = re.sub(r"\.[a-z0-9]+$", "", name.lower())
        return re.sub(r"[^a-z0-9]+", " ", base).strip()

    query_lower = query.lower()
    multi_file_intent = any(
        phrase in query_lower
        for phrase in [
            "both files",
            "all files",
            "compare files",
            "across files",
            "from all documents",
            "combine",
            "summarize all",
        ]
    )

    scored = [(score_chunk(query, item["text"]), item) for item in data]
    scored = [item for item in scored if item[0] > 0]
    scored.sort(key=lambda entry: entry[0], reverse=True)

    # If user asked about a specific file name, scope retrieval to that file.
    available_files = sorted({item["file_name"] for item in data})
    explicit_targets: List[str] = []
    for file_name in available_files:
        file_name_lower = file_name.lower()
        normalized_token = normalize_file_token(file_name)
        if file_name_lower in query_lower or (normalized_token and normalized_token in query_lower):
            explicit_targets.append(file_name)

    if explicit_targets:
        scored = [entry for entry in scored if entry[1]["file_name"] in explicit_targets]

    # Keep cross-file retrieval enabled by default so multi-file questions
    # can be answered from more than one document.
    # If explicit file targets are present, retrieval is already scoped above.

    # Group candidates by file so answers can synthesize across multiple files.
    by_file: Dict[str, List[Dict]] = {}
    for _, item in scored:
        by_file.setdefault(item["file_name"], []).append(item)

    selected: List[Dict] = []
    seen_keys = set()
    file_counts: Dict[str, int] = {}
    file_names = sorted(by_file.keys())

    # Round-robin pick: first pass prioritizes breadth (one chunk per file),
    # then fills remaining slots while respecting per-file cap.
    made_progress = True
    while made_progress and len(selected) < limit:
        made_progress = False
        for file_name in file_names:
            if len(selected) >= limit:
                break
            if file_counts.get(file_name, 0) >= MAX_CHUNKS_PER_FILE:
                continue
            candidates = by_file.get(file_name, [])
            if not candidates:
                continue
            item = candidates.pop(0)
            key = re.sub(r"\s+", " ", item.get("text", "")[:220]).strip().lower()
            if not key or key in seen_keys:
                continue
            seen_keys.add(key)
            selected.append(item)
            file_counts[file_name] = file_counts.get(file_name, 0) + 1
            made_progress = True

    # If no lexical hits, include one chunk per uploaded file as a safe fallback.
    if not selected and data:
        fallback_by_file: Dict[str, Dict] = {}
        for item in data:
            if item["file_name"] not in fallback_by_file:
                fallback_by_file[item["file_name"]] = item
        selected = list(fallback_by_file.values())[:limit]

    return selected


def build_history_context(messages: List[Dict], max_items: int = 6) -> str:
    valid = [item for item in messages if item.get("role") in {"user", "assistant"}]
    recent = valid[-max_items:]
    lines: List[str] = []
    for item in recent:
        role = item.get("role", "user").upper()
        text = re.sub(r"\s+", " ", str(item.get("content", ""))).strip()
        if text:
            lines.append(f"{role}: {text[:500]}")
    return "\n".join(lines)


def build_prompt(question: str, context_chunks: List[Dict], style: str, history_text: str) -> str:
    if style not in ALLOWED_ANSWER_STYLES:
        style = "Detailed"
    style_prompt = {
        "Detailed": (
            "Respond in clear plain language with complete explanation.\n"
            "Use short paragraphs and optional bullet points.\n"
            "Do not use section heading labels."
        ),
        "Short": "Respond in 2-4 concise sentences only.",
        "Bullet Points": "Respond only in bullet points with 5-8 bullets.",
        "Professional": (
            "Respond in a professional tone with clear, well-ordered paragraphs.\n"
            "Use bullets only where needed.\n"
            "Do not use heading labels like Executive Summary, Findings, or Recommendation."
        ),
    }[style]

    if context_chunks:
        context_lines = []
        for chunk in context_chunks:
            context_lines.append(
                f"File: {chunk['file_name']} | {chunk['location']}\n{chunk['text']}"
            )
        context_text = "\n\n".join(context_lines)
        file_count = len({chunk["file_name"] for chunk in context_chunks})
        if file_count > 1:
            context_instruction = (
                "Use the context below as the primary evidence. "
                "Synthesize across multiple files when relevant instead of answering from only one file. "
                "When multiple files are used, clearly tag each key point with its source file name."
            )
        else:
            context_instruction = "Use the context below as the primary evidence for your answer."
    else:
        context_text = "No document context available."
        context_instruction = "No source context is available, so answer directly."

    return f"""
You are a reliable, highly organized assistant.
{style_prompt}
{context_instruction}
Follow only the selected answer style.
If context is missing important details, clearly say what is missing.
Do not include raw source tags like [S1] or mention internal prompt instructions.
Do not repeat the same sentence or point.
When exact wording exists in context, use that wording once and keep the answer concise and precise.
Every line must be a complete sentence with proper grammar and punctuation.
Avoid sentence fragments, token lists, repeated clauses, and malformed text.

Question:
{question}

Recent conversation:
{history_text or "No prior conversation context."}

Context:
{context_text}
""".strip()


def is_low_quality_fragment(text: str) -> bool:
    value = re.sub(r"\s+", " ", (text or "")).strip()
    if not value:
        return True
    if len(value) < 22:
        return True
    alpha = sum(1 for ch in value if ch.isalpha())
    digits = sum(1 for ch in value if ch.isdigit())
    if alpha == 0:
        return True
    digit_ratio = digits / max(len(value), 1)
    if digit_ratio > 0.35:
        return True
    words = re.findall(r"[A-Za-z]{2,}", value)
    if len(words) < 4:
        return True
    upper_words = [w for w in words if w.isupper() and len(w) > 2]
    if upper_words and len(upper_words) / len(words) > 0.65:
        return True
    return False


def format_answer_by_style(answer: str, style: str) -> str:
    text = (answer or "").strip()
    if not text:
        return text

    def normalize_sentence_case(line: str) -> str:
        if not line:
            return line
        if line[0].isalpha() and line[0].islower():
            return line[0].upper() + line[1:]
        return line

    def clean_line(line: str) -> str:
        line = re.sub(r"\s+", " ", line or "").strip()
        if not line:
            return ""
        if line.startswith("#"):
            return line
        line = re.sub(r"([,;:])(?=\S)", r"\1 ", line)
        line = re.sub(r"\s+([,.!?])", r"\1", line)
        line = re.sub(r"([.!?]){2,}", r"\1", line)
        line = normalize_sentence_case(line)
        if line.startswith("- "):
            body = line[2:].strip()
            if is_low_quality_fragment(body):
                return ""
            if body and body[-1] not in ".!?":
                body += "."
            return f"- {body}"
        if is_low_quality_fragment(line):
            return ""
        if line and line[-1] not in ".!?:":
            line += "."
        return line

    def semantic_key(text_line: str) -> str:
        base = re.sub(r"[^a-z0-9 ]", " ", text_line.lower())
        tokens = [tok for tok in base.split() if len(tok) > 2]
        return " ".join(tokens[:18])

    def too_similar(a: str, b: str) -> bool:
        a_tokens = set(semantic_key(a).split())
        b_tokens = set(semantic_key(b).split())
        if not a_tokens or not b_tokens:
            return False
        overlap = len(a_tokens & b_tokens) / max(len(a_tokens | b_tokens), 1)
        return overlap >= 0.82

    def dedupe_lines(multiline_text: str) -> str:
        lines = [ln for ln in multiline_text.splitlines() if ln.strip()]
        output: List[str] = []
        seen = set()
        for line in lines:
            cleaned = clean_line(line)
            if not cleaned:
                continue
            key = re.sub(r"[^a-z0-9 ]", "", cleaned.lower()).strip()
            if key and key in seen:
                continue
            if any(too_similar(cleaned, existing) for existing in output):
                continue
            if key:
                seen.add(key)
            output.append(cleaned)
        return "\n".join(output).strip()

    def dedupe_sentences(single_text: str) -> str:
        parts = [p.strip() for p in re.split(r"(?<=[.!?])\s+", single_text) if p.strip()]
        output: List[str] = []
        seen = set()
        for part in parts:
            fixed = clean_line(part)
            if not fixed:
                continue
            key = re.sub(r"[^a-z0-9 ]", "", fixed.lower()).strip()
            if key and key in seen:
                continue
            if any(too_similar(fixed, existing) for existing in output):
                continue
            if key:
                seen.add(key)
            output.append(fixed)
        return " ".join(output).strip()

    text = re.sub(r"\[S\d+\]", "", text)
    text = re.sub(r"^##\s*(Executive Summary|Findings|Risks/Gaps|Recommendation)\s*$", "", text, flags=re.I | re.M)
    text = re.sub(r"^(Executive Summary|Findings|Risks/Gaps|Recommendation)\s*:?\s*$", "", text, flags=re.I | re.M)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    text = re.sub(r"\s{2,}", " ", text)

    if style == "Short":
        sentences = re.split(r"(?<=[.!?])\s+", text)
        trimmed = " ".join([s for s in sentences if s][:4]).strip()
        return dedupe_sentences(trimmed or text[:350])

    if style == "Bullet Points":
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        bullets = [ln for ln in lines if ln.startswith(("-", "*"))]
        if bullets:
            return dedupe_lines("\n".join(bullets[:8]))
        chunks = re.split(r"(?<=[.!?])\s+", text)
        chunks = [c.strip() for c in chunks if c.strip()]
        return dedupe_lines("\n".join([f"- {item}" for item in chunks[:8]]))

    if style == "Professional":
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        return dedupe_lines("\n".join(lines))

    if style == "Detailed":
        return dedupe_lines(text)

    return dedupe_sentences(text)


def is_api_key_issue(error_text: str) -> bool:
    lowered = (error_text or "").lower()
    return any(pattern.lower() in lowered for pattern in API_KEY_ERROR_PATTERNS)


def _query_terms(question: str) -> List[str]:
    terms = [word.strip(".,!?;:\"'()[]{}").lower() for word in question.split()]
    base = [word for word in terms if len(word) > 2]
    phrase_map = {
        "problem statement": "problem statement",
        "main objective": "main objective",
        "key challenges": "key challenges",
        "flow chart": "flow chart",
        "flowchart": "flowchart",
        "diagram": "diagram",
        "table": "table",
    }
    lower_q = question.lower()
    for _, phrase in phrase_map.items():
        if phrase in lower_q and phrase not in base:
            base.append(phrase)
    return base


def emphasize_answer_focus(answer: str, question: str, style: str) -> str:
    # Disabled intentionally to avoid repetitive highlighted prefixes in final answers.
    return (answer or "").strip()


def _extract_relevant_sentences(question: str, context_chunks: List[Dict], limit: int = 6) -> List[str]:
    terms = _query_terms(question)
    candidates: List[tuple[int, str]] = []

    for chunk in context_chunks[:8]:
        raw = chunk.get("text", "") or ""
        raw = raw.replace("|", ", ")
        raw = re.sub(r"\s{2,}", ", ", raw)
        raw = re.sub(r"\s+", " ", raw).strip()
        if not raw:
            continue
        sentences = re.split(r"(?<=[.!?])\s+", raw)
        for sentence in sentences:
            sentence = sentence.strip()
            if is_low_quality_fragment(sentence):
                continue
            lowered = sentence.lower()
            score = sum(3 for term in terms if term in lowered)
            if "problem statement" in lowered:
                score += 5
            if score > 0:
                candidates.append((score, sentence))

    candidates.sort(key=lambda item: item[0], reverse=True)
    chosen: List[str] = []
    seen = set()
    for _, sentence in candidates:
        norm = re.sub(r"[^a-z0-9 ]", "", sentence.lower()).strip()
        key = " ".join(norm.split()[:20])
        if key in seen:
            continue
        too_similar = any(key in existing or existing in key for existing in seen if existing)
        if too_similar:
            continue
        seen.add(key)
        chosen.append(sentence)
        if len(chosen) >= limit:
            break

    if chosen:
        return chosen

    # fallback if lexical match is weak
    defaults: List[str] = []
    for chunk in context_chunks[:3]:
        raw = re.sub(r"\s+", " ", chunk.get("text", "")).strip()
        if not raw:
            continue
        first = re.split(r"(?<=[.!?])\s+", raw)[0].strip()
        if not is_low_quality_fragment(first):
            defaults.append(first)
    return defaults[:limit]


def fallback_answer_from_context(question: str, context_chunks: List[Dict], style: str) -> str:
    if not context_chunks:
        if style == "Detailed":
            return (
                "I cannot generate a full detailed response yet because no document context is currently indexed. "
                "Please upload files and wait for auto-processing. Once indexing is complete, I will provide a full "
                "detailed answer with explanation, key observations, and structured takeaways."
            )
        return (
            "I need either a working API key for general answers or processed files for RAG-based answers. "
            "Please upload files so I can retrieve context and answer from your documents."
        )

    snippets = _extract_relevant_sentences(question, context_chunks, limit=4)
    if not snippets:
        snippets = ["The uploaded context contains limited directly matchable details for this question."]

    if style == "Short":
        return snippets[0]

    if style == "Bullet Points":
        lines = ["Most relevant points from your uploaded documents:"]
        for idx, snippet in enumerate(snippets, start=1):
            lines.append(f"- Point {idx}: {snippet}")
        return "\n".join(lines)

    if style == "Professional":
        lines = [snippets[0], "", "Key supporting points:"]
        for item in snippets[1:4]:
            lines.append(f"- {item}")
        lines.extend(["", "This response is grounded in currently processed document context."])
        return "\n".join(lines)

    lines = [snippets[0], "", "Key points:"]
    for item in snippets[1:4]:
        lines.append(f"- {item}")
    lines.extend(["", "This answer is generated from processed files currently available in context."])
    return "\n".join(lines)


def source_files_label(sources: List[Dict]) -> str:
    if not sources:
        return "General response (no specific file matched)."
    names = []
    seen = set()
    for item in sources:
        file_name = str(item.get("file_name", "")).strip()
        if file_name and file_name not in seen:
            seen.add(file_name)
            names.append(file_name)
    if not names:
        return "General response (no specific file matched)."
    return ", ".join(names)


def stream_model_response(model, prompt: str, placeholder) -> tuple[str, bool, str]:
    for attempt in range(1, MAX_API_RETRIES + 1):
        try:
            fragments: List[str] = []
            if model.get("backend") == "google_genai":
                response_stream = model["client"].models.generate_content_stream(
                    model=model["model_name"],
                    contents=prompt,
                    config={"temperature": 0.25, "top_p": 0.9, "max_output_tokens": 1400},
                )
            else:
                response_stream = model["model"].generate_content(
                    prompt,
                    stream=True,
                    generation_config={"temperature": 0.25, "top_p": 0.9, "max_output_tokens": 1400},
                )
            for piece in response_stream:
                text = getattr(piece, "text", "")
                if text:
                    fragments.append(text)
                    placeholder.markdown("".join(fragments) + "▌")

            answer = "".join(fragments).strip()
            if answer:
                return answer, True, ""
            break
        except Exception as exc:
            if attempt == MAX_API_RETRIES:
                return "", False, str(exc)
            time.sleep(min(0.8 * attempt, 2.2))

    return "", False, "No response text returned by API."


def call_model_non_stream(model, prompt: str) -> tuple[str, bool, str]:
    for attempt in range(1, MAX_API_RETRIES + 1):
        try:
            if model.get("backend") == "google_genai":
                response = model["client"].models.generate_content(
                    model=model["model_name"],
                    contents=prompt,
                    config={"temperature": 0.2, "top_p": 0.9, "max_output_tokens": 1400},
                )
            else:
                response = model["model"].generate_content(
                    prompt,
                    generation_config={"temperature": 0.2, "top_p": 0.9, "max_output_tokens": 1400},
                )
            answer = (getattr(response, "text", "") or "").strip()
            if answer:
                return answer, True, ""
            return "", False, "No response text returned by API."
        except Exception as exc:
            if attempt == MAX_API_RETRIES:
                return "", False, str(exc)
            time.sleep(min(0.8 * attempt, 2.2))
    return "", False, "Non-stream API call failed."


def build_compact_prompt(question: str, context_chunks: List[Dict], style: str) -> str:
    compact_chunks = context_chunks[:3]
    compact_context = "\n\n".join(
        [f"File: {item['file_name']} | {item['location']}\n{item['text'][:450]}" for item in compact_chunks]
    )
    if not compact_context:
        compact_context = "No document context available."

    return f"""
You are a precise assistant. Follow only this style: {style}.
Answer clearly with complete, non-repetitive sentences and proper punctuation.
Do not use heading labels.
Use only the relevant details from context.

Question:
{question}

Context:
{compact_context}
""".strip()


def generate_answer_with_recovery(
    model,
    prompt: str,
    placeholder,
    question: str,
    context_chunks: List[Dict],
    style: str,
) -> tuple[str, bool, str]:
    answer, ok, error_text = stream_model_response(model, prompt, placeholder)
    if ok:
        return answer, True, ""

    answer, ok, non_stream_error = call_model_non_stream(model, prompt)
    if ok:
        return answer, True, ""

    compact_prompt = build_compact_prompt(question, context_chunks, style)
    answer, ok, compact_error = call_model_non_stream(model, compact_prompt)
    if ok:
        return answer, True, ""

    merged_error = " | ".join([msg for msg in [error_text, non_stream_error, compact_error] if msg])
    return "", False, merged_error or "All API generation attempts failed."


def merge_uploaded_chunks(existing: List[Dict], incoming: List[Dict]) -> List[Dict]:
    incoming_files = {chunk["file_name"] for chunk in incoming}
    retained = [chunk for chunk in existing if chunk["file_name"] not in incoming_files]
    return retained + incoming


def get_upload_signature(uploaded_files) -> tuple[str, ...]:
    if not uploaded_files:
        return tuple()
    signature = []
    for file in uploaded_files:
        size = getattr(file, "size", 0)
        signature.append(f"{file.name}:{size}")
    return tuple(sorted(signature))


def is_mobile_client() -> bool:
    try:
        user_agent = str(st.context.headers.get("user-agent", "")).lower()
    except Exception:
        user_agent = ""

    mobile_markers = [
        "android",
        "iphone",
        "ipad",
        "ipod",
        "mobile",
        "opera mini",
        "windows phone",
    ]
    return any(marker in user_agent for marker in mobile_markers)


def process_uploaded_files(uploaded_files) -> tuple[int, int]:
    incoming_chunks: List[Dict] = []
    processed = 0
    failed = 0

    for file in uploaded_files or []:
        file.seek(0)
        chunks = load_uploaded_file(file)
        if chunks:
            incoming_chunks.extend(chunks)
            processed += 1
            st.session_state.file_status[file.name] = "processed"
        else:
            failed += 1
            st.session_state.file_status[file.name] = "failed"

    if incoming_chunks:
        st.session_state.data = merge_uploaded_chunks(st.session_state.data, incoming_chunks)
        names = sorted({item["file_name"] for item in st.session_state.data})
        st.session_state.uploaded_file_names = names
        st.session_state.process_batches += 1
    return processed, failed


def render_runtime_status(slot, phase: str, detail: str, done: bool = False) -> None:
    status_class = "runtime-status done" if done else "runtime-status"
    slot.markdown(
        f"""
<div class="{status_class}">
  <div class="runtime-dot"></div>
  <strong>{phase}</strong>
  <span>{detail}</span>
</div>
""",
        unsafe_allow_html=True,
    )


def render_sidebar_rag_pipeline(model_connected: bool) -> None:
    total_files = len(st.session_state.uploaded_file_names)
    total_chunks = len(st.session_state.data)
    total_questions = sum(1 for item in st.session_state.messages if item.get("role") == "user")
    ready_for_retrieval = total_chunks > 0 and total_files > 0
    avg_chunks_per_file = round(total_chunks / total_files, 1) if total_files else 0
    complete_steps = sum([1 if total_files else 0, 1 if total_chunks else 0, 1 if ready_for_retrieval else 0, 1 if model_connected else 0])
    progress_pct = int((complete_steps / 4) * 100)

    ingest_class = "ok" if total_files else "wait"
    chunk_class = "ok" if total_chunks else "wait"
    retrieve_class = "ok" if ready_for_retrieval else "wait"
    generate_class = "ok" if model_connected else "off"

    st.sidebar.markdown("### RAG Quick Flow")
    st.sidebar.markdown(
        f"""
<div class="rag-lite">
  <div class="rag-lite-head">Retrieval Pipeline</div>
  <div class="rag-lite-sub">Fast status view for your chatbot flow</div>

  <div class="rag-flow-row">
    <div class="rag-stage">1) Ingest</div>
    <div class="rag-pill {ingest_class}">{'READY' if total_files else 'WAIT'}</div>
  </div>
  <div class="rag-flow-row">
    <div class="rag-stage">2) Index</div>
    <div class="rag-pill {chunk_class}">{'READY' if total_chunks else 'WAIT'}</div>
  </div>
  <div class="rag-flow-row">
    <div class="rag-stage">3) Retrieve</div>
    <div class="rag-pill {retrieve_class}">{'ACTIVE' if ready_for_retrieval else 'STANDBY'}</div>
  </div>
  <div class="rag-flow-row">
    <div class="rag-stage">4) Generate</div>
    <div class="rag-pill {generate_class}">{'ONLINE' if model_connected else 'OFFLINE'}</div>
  </div>

  <div class="rag-progress"><div style="width:{progress_pct}%"></div></div>
  <div class="rag-mini">
    Progress: {progress_pct}% | Files: {total_files} | Chunks: {total_chunks}<br/>
    Avg Chunks/File: {avg_chunks_per_file} | Queries: {total_questions}
  </div>
</div>
""",
        unsafe_allow_html=True,
    )


def render_sidebar_dashboard(model_connected: bool) -> None:
    messages = st.session_state.messages
    total_questions = sum(1 for item in messages if item.get("role") == "user")
    total_answers = sum(1 for item in messages if item.get("role") == "assistant")
    total_files = len(st.session_state.uploaded_file_names)
    total_chunks = len(st.session_state.data)
    total_citations = sum(len(item.get("sources", [])) for item in messages if item.get("role") == "assistant")
    api_failures = st.session_state.api_failures
    processed_batches = st.session_state.process_batches

    st.sidebar.markdown("### Numeric Dashboard")
    top_row = st.sidebar.columns(3)
    top_row[0].metric("API Connected", 1 if model_connected else 0)
    top_row[1].metric("Questions Asked", total_questions)
    top_row[2].metric("Answers Given", total_answers)

    extra_row = st.sidebar.columns(2)
    extra_row[0].metric("Files Loaded", total_files)
    extra_row[1].metric("Chunks Indexed", total_chunks)

    extra_row_2 = st.sidebar.columns(2)
    extra_row_2[0].metric("Source Citations", total_citations)
    extra_row_2[1].metric("Process Runs", processed_batches)

    st.sidebar.metric("API Failure Events", api_failures)


def init_session_state() -> None:
    if "messages" not in st.session_state:
        st.session_state.messages = load_chat()
    if "data" not in st.session_state:
        st.session_state.data = []
    if "uploaded_file_names" not in st.session_state:
        st.session_state.uploaded_file_names = []
    if "api_failures" not in st.session_state:
        st.session_state.api_failures = 0
    if "api_consecutive_failures" not in st.session_state:
        st.session_state.api_consecutive_failures = 0
    if "api_last_request_ok" not in st.session_state:
        st.session_state.api_last_request_ok = True
    if "api_key_invalid" not in st.session_state:
        st.session_state.api_key_invalid = False
    if "api_failure_event_open" not in st.session_state:
        st.session_state.api_failure_event_open = False
    if "process_batches" not in st.session_state:
        st.session_state.process_batches = 0
    if "file_status" not in st.session_state:
        st.session_state.file_status = {}
    if "last_upload_signature" not in st.session_state:
        st.session_state.last_upload_signature = tuple()
    if "mobile_uploader_nonce" not in st.session_state:
        st.session_state.mobile_uploader_nonce = 0


def main() -> None:
    st.set_page_config(
        page_title=PROJECT_TITLE,
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    apply_ui_styles()
    init_session_state()

    st.markdown(
        f"""
<div class="app-header">
  <h1>{PROJECT_TITLE}</h1>
  <p>Next-gen animated interface with instant multi-file indexing, retrieval, and grounded conversational responses.</p>
</div>
""",
        unsafe_allow_html=True,
    )

    api_key = os.getenv("GEMINI_API_KEY", "").strip()
    model = configure_model(api_key)
    if model:
        st.session_state.api_key_invalid = False
    api_available = bool(model)
    model_connected = api_available
    if model_connected:
        st.success("API status: Connected")
    else:
        st.warning("API status: Disconnected (file processing still works without API)")

    st.markdown('<div class="upload-panel">', unsafe_allow_html=True)
    control_left, control_right = st.columns([1, 2])
    mobile_client = is_mobile_client()
    with control_left:
        answer_style = st.selectbox(
            "Answer style",
            ALLOWED_ANSWER_STYLES,
            help="Choose how the assistant should format responses.",
        )
    with control_right:
        if mobile_client:
            st.caption("Mobile upload mode: upload one file at a time for best reliability.")
            st.caption("Chrome tip: pick files from Downloads/Files, and keep filename in English.")
            uploaded_single = st.file_uploader(
                "📂 Tap to browse files",
                type=None,
                accept_multiple_files=False,
                help="Mobile tip: use English file names and upload one file at a time.",
                key=f"mobile_file_uploader_{st.session_state.mobile_uploader_nonce}",
            )
            if uploaded_single:
                ext = Path(uploaded_single.name).suffix.lower().replace(".", "").strip()
                if ext not in SUPPORTED_TYPES:
                    st.error(
                        f"Unsupported file type on mobile: .{ext or 'unknown'}. "
                        f"Allowed: {', '.join(SUPPORTED_TYPES).upper()}"
                    )
                    uploaded_files = []
                else:
                    uploaded_files = [uploaded_single]
            else:
                uploaded_files = []
        else:
            uploaded_files = st.file_uploader(
                "📂 Drag and drop files here or browse files",
                type=SUPPORTED_TYPES,
                accept_multiple_files=True,
                help="Supported: PDF, TXT, XML, DOCX, JSON, CSV",
                key="desktop_file_uploader",
            )

    upload_signature = get_upload_signature(uploaded_files)
    if upload_signature and upload_signature != st.session_state.last_upload_signature:
        with st.spinner("Processing uploaded files..."):
            processed, failed = process_uploaded_files(uploaded_files)
        st.session_state.last_upload_signature = upload_signature
        if processed or failed:
            st.success(
                f"Auto-processed {processed} file(s), failed {failed}. Context: {len(st.session_state.uploaded_file_names)} files / {len(st.session_state.data)} chunks."
            )
        # Reset mobile uploader after each processing cycle to avoid
        # stuck/error UI state in some Android Chrome sessions.
        if mobile_client and processed > 0:
            st.session_state.mobile_uploader_nonce += 1
            st.rerun()

    if uploaded_files:
        selected_names = [file.name for file in uploaded_files]
        st.caption(f"Selected files: {len(selected_names)}")
        status_icon = {"processed": "✅", "failed": "❌"}
        status_label = {"processed": "Processed", "failed": "Failed"}
        for name in selected_names:
            state = st.session_state.file_status.get(name)
            if state in status_icon:
                st.caption(f"{status_icon[state]} {name} - {status_label[state]}")
    else:
        st.session_state.last_upload_signature = tuple()

    action_1, action_2 = st.columns(2)
    clear_clicked = action_1.button("🧹 Clear Chat", use_container_width=True)
    reset_clicked = action_2.button("🗑️ Reset Documents", use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

    if clear_clicked:
        st.session_state.messages = []
        save_chat([])
        st.success("Chat history cleared.")

    if reset_clicked:
        st.session_state.data = []
        st.session_state.uploaded_file_names = []
        st.session_state.file_status = {}
        st.session_state.last_upload_signature = tuple()
        st.success("Document context reset.")

    render_sidebar_rag_pipeline(model_connected=model_connected)
    render_sidebar_dashboard(model_connected=model_connected)

    if st.session_state.uploaded_file_names:
        st.markdown("#### Processed Files")
        per_file_chunks: Dict[str, int] = {}
        for item in st.session_state.data:
            per_file_chunks[item["file_name"]] = per_file_chunks.get(item["file_name"], 0) + 1
        for name in st.session_state.uploaded_file_names:
            st.caption(f"📄 {name}  •  {per_file_chunks.get(name, 0)} chunks")

    for message in st.session_state.messages:
        with st.chat_message(message.get("role", "assistant")):
            st.markdown(message.get("content", ""))
            if message.get("role") == "assistant":
                source_line = source_files_label(message.get("sources", []))
                st.caption(f"Sources used: {source_line}")

    question = st.chat_input("Ask anything about your files...")
    if not question:
        return

    user_message = {"role": "user", "content": question}
    st.session_state.messages.append(user_message)
    save_chat(st.session_state.messages)

    with st.chat_message("user"):
        st.markdown(question)

    relevant_chunks = top_relevant_chunks(question, st.session_state.data)
    history_text = build_history_context(st.session_state.messages[:-1])
    prompt = build_prompt(question, relevant_chunks, answer_style, history_text)

    with st.chat_message("assistant"):
        status_slot = st.empty()
        placeholder = st.empty()
        render_runtime_status(
            status_slot,
            "Analyzing Request",
            "Understanding the question and preparing the response flow.",
        )
        render_runtime_status(
            status_slot,
            "Retrieving Context",
            "Collecting the most relevant document evidence for this query.",
        )
        if api_available:
            render_runtime_status(
                status_slot,
                "Generating Response",
                "Composing a grounded response from the retrieved context.",
            )
            answer, ok, error_text = generate_answer_with_recovery(
                model=model,
                prompt=prompt,
                placeholder=placeholder,
                question=question,
                context_chunks=relevant_chunks,
                style=answer_style,
            )
            if not ok:
                st.session_state.api_consecutive_failures += 1
                st.session_state.api_last_request_ok = False
                if not st.session_state.api_failure_event_open:
                    st.session_state.api_failures += 1
                    st.session_state.api_failure_event_open = True
                if is_api_key_issue(error_text):
                    st.session_state.api_key_invalid = True
                render_runtime_status(
                    status_slot,
                    "Local Recovery Mode",
                    "Live API response was unavailable, switching to document-grounded fallback.",
                )
                answer = fallback_answer_from_context(question, relevant_chunks, answer_style)
            else:
                st.session_state.api_last_request_ok = True
                st.session_state.api_consecutive_failures = 0
                st.session_state.api_failure_event_open = False
        else:
            st.session_state.api_last_request_ok = False
            render_runtime_status(
                status_slot,
                "Local Document Mode",
                "API is offline, generating response from processed files only.",
            )
            answer = fallback_answer_from_context(question, relevant_chunks, answer_style)

        if not answer.strip():
            answer = fallback_answer_from_context(question, relevant_chunks, answer_style)
        answer = format_answer_by_style(answer, answer_style)
        if not answer.strip():
            answer = (
                "I could not construct a clean response from the current text extraction. "
                "Please upload a clearer source file or ask a narrower question."
            )
        placeholder.markdown(answer)
        st.caption(f"Sources used: {source_files_label(relevant_chunks)}")
        render_runtime_status(
            status_slot,
            "Response Ready",
            "Answer generated successfully and delivered to chat.",
            done=True,
        )

    assistant_message = {
        "role": "assistant",
        "content": answer,
        "style": answer_style,
        "sources": relevant_chunks,
    }
    st.session_state.messages.append(assistant_message)
    save_chat(st.session_state.messages)


if __name__ == "__main__":
    main()
